import docx 
from collections import Counter
import re # Модуль для работы с текстом 
import pandas as pd
import matplotlib.pyplot as plt

doc = docx.Document('lion.docx') 
num_par = len(doc.paragraphs) # Количество абзацев
word_list = [] # Список слов без повторений 

a = []
rus_let = re.compile(r'[а-яё]', re.IGNORECASE) # Выражение для поиска русских букв

rus_text = [] # Список для всего текста (русский)
rus_word = re.compile(r'\b[а-яё]+\b',re.IGNORECASE) # Выражение для поиска русских слов 

for paragraph in doc.paragraphs: # Проходимся по всем параграфам в документе
    text = paragraph.text.lower() # Приводим текст к нижнему регистру
    words = rus_word.findall(text) # Находим все русские слова
    letters = rus_let.findall(text) # Находим все русские буквы
    a.extend(letters) # Добовляем буквы в общий список

    for word in words:
        rus_text.append(word)
        if word not in word_list: # Проверка на наличие слов в списке 
            word_list.append(word) # Добавления слова 

# Слова в тексте
rus_text_len = len(rus_text) # Количество слов в списке

word_counts = Counter(rus_text) # Подсчитываем частоту встречаемости слов

dict_word = {
    'Слово': list(word_counts.keys()),
    'Частота встречи,раз': list(word_counts.values()),
    'Частота встречи в %': [round((count / rus_text_len) * 100, 2) for count in word_counts.values()]
}

df_word = pd.DataFrame(dict_word)

doc_word = docx.Document() # Создаю новый документ

doc_word.add_heading('Встречаемость слов в тексте', 0) # Добовляем заголовок

table = doc_word.add_table(rows=1, cols=3) # Создаем таблицу с одной строкой и тремя колонками 
hdr_cells = table.rows[0].cells # cells - используем для доступа к конкретным ячейкам строки, чтобы заполнить их содержимым
hdr_cells[0].text = 'Слово'
hdr_cells[1].text = 'Частота встречи,раз'
hdr_cells[2].text = 'Частота встречи в %'

for i in range(len(df_word)):
    row_cells = table.add_row().cells # Добавляем новой строки в таблицу                                        
    row_cells[0].text = df_word.loc[i, 'Слово'] # loc - выбор строки с индексом [i]
    row_cells[1].text = str(df_word.loc[i, 'Частота встречи,раз'])
    row_cells[2].text = str(df_word.loc[i, 'Частота встречи в %'])

doc_word.save('встречаемость_слов.docx') # Сохраняем документ

# Буквы в тексте

letters_counts = Counter(a) # Подсчитываем частоту встречаемости букв
df_letter = pd.DataFrame(letters_counts.items(),columns=['Буквы','Количество'])
df_letter['Количество'] = pd.to_numeric(df_letter['Количество']) # to_numeric - преобразования значений в столбцах DataFrame в числовой тип данных

plt.figure(figsize=(10, 6))# Ширина фигуры будет 10 дюймов, а высота — 6 дюймов.
plt.bar(df_letter['Буквы'], df_letter['Количество'], color='skyblue') # plt.bar - создаем столбчатую диаграмму.
plt.xlabel('Буквы') # оси X, представляющие буквы
plt.ylabel('Количество') # оси Y, представляющие Количество
plt.title('Встречаемость букв в тексте') # plt.title - задаем заголовок для графика.
plt.grid(True) # plt.grid - включаем отображение сетки на графике, что помогает лучше визуализировать данные.
plt.show()

